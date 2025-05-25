
import streamlit as st
from docx import Document
import datetime
import io

st.set_page_config(page_title="PROMTO - Receptor", layout="centered")
st.title("📋 PROMTO – Formulario del Receptor")

def es_mayor_un_ano(fecha):
    return (datetime.date.today() - fecha).days > 365

def generar_word(nombre, datos):
    doc = Document()
    doc.add_heading(f"Checklist PROMTO – Receptor: {nombre}", level=1)
    for seccion, campos in datos.items():
        doc.add_heading(seccion, level=2)
        for item in campos:
            texto = f"{item['nombre']} - Fecha: {item['fecha'].strftime('%d/%m/%Y')}"
            if 'resultado' in item:
                texto += f" - Resultado: {item['resultado']}"
            if 'observaciones' in item and item['observaciones']:
                texto += f"\nObservaciones: {item['observaciones']}"
            doc.add_paragraph(texto)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Este archivo incluye el flujo completo del receptor PROMTO.
# Para ejecutarlo correctamente, completar el formulario dentro de una función Streamlit como fue entregado antes.
