
import streamlit as st
import datetime
from docx import Document
import io

st.set_page_config(page_title="PROMTO - Receptor", layout="centered")
st.title("🧾 PROMTO - Formulario Receptor")

# Función de validación de fecha > 1 año
def es_mayor_un_ano(fecha):
    return (datetime.date.today() - fecha).days > 365

# Función para generar el documento Word
def generar_word(data):
    doc = Document()
    doc.add_heading("Checklist PROMTO - Receptor", level=1)
    for seccion, campos in data.items():
        doc.add_heading(seccion, level=2)
        for item in campos:
            texto = f"{item['nombre']} - Fecha: {item['fecha'].strftime('%d/%m/%Y')}"
            if 'resultado' in item:
                texto += f" - Resultado: {item['resultado']}"
            if 'observaciones' in item and item['observaciones']:
                texto += f"\nObservaciones: {item['observaciones']}"
            doc.add_paragraph(texto)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ---- FORMULARIO ----
with st.form("form_receptor"):
    datos = {}

    # Ecocardiograma
    st.subheader("🫀 Ecocardiograma")
    fecha_eco = st.date_input("Fecha", key="eco_fecha")
    fevi = st.text_input("FEVI (%)", key="fevi")
    psap = st.text_input("PSAP (mmHg)", key="psap")
    valvulopatia = st.selectbox("¿Valvulopatías u otras alteraciones?", ["No", "Sí"], key="valv")
    observaciones_valv = ""
    if valvulopatia == "Sí":
        observaciones_valv = st.text_area("Observaciones", key="obs_valv")

    datos["Ecocardiograma"] = [{
        "nombre": "Ecocardiograma",
        "fecha": fecha_eco,
        "resultado": f"FEVI: {fevi}, PSAP: {psap}, Valvulopatía: {valvulopatia}",
        "observaciones": observaciones_valv
    }]

    # Estudios de imagen
    st.subheader("🖼️ Estudios de imagen")
    imagenes = {}
    for estudio in ["Tomografía de abdomen", "Tomografía de senos paranasales", "Radiografía/TAC de tórax"]:
        tiene_hallazgos = st.selectbox(f"¿{estudio} con hallazgos patológicos?", ["No", "Sí"], key=estudio+"_hall")
        dictado = st.selectbox(f"¿{estudio} tiene dictado oficial?", ["No", "Sí"], key=estudio+"_dict")
        fecha = st.date_input(f"Fecha de {estudio}", key=estudio+"_fecha")
        obs = ""
        if tiene_hallazgos == "Sí":
            obs = st.text_area(f"Observaciones de {estudio}", key=estudio+"_obs")
        imagenes[estudio] = [{
            "nombre": estudio,
            "fecha": fecha,
            "resultado": f"Hallazgos: {tiene_hallazgos}, Dictado: {dictado}",
            "observaciones": obs
        }]
    datos["Estudios de Imagen"] = sum(imagenes.values(), [])

    # Estudios de microbiología
    st.subheader("🦠 Microbiología")
    micro = {}
    for estudio in ["Exudado nasal", "Exudado faríngeo", "Quantiferon", "VDRL", "Urocultivo"]:
        fecha = st.date_input(f"Fecha de {estudio}", key=estudio+"_fecha")
        resultado = st.selectbox(f"Resultado de {estudio}", ["Negativo", "Positivo"], key=estudio+"_res")
        obs = ""
        if resultado == "Positivo":
            obs = st.text_area(f"Observaciones de {estudio}", key=estudio+"_obs")
        color = "🔴 " if es_mayor_un_ano(fecha) else ""
        micro[estudio] = [{
            "nombre": f"{color}{estudio}",
            "fecha": fecha,
            "resultado": resultado,
            "observaciones": obs
        }]
    datos["Microbiología"] = sum(micro.values(), [])

    # Enviar
    enviado = st.form_submit_button("✅ Generar Checklist")

if enviado:
    archivo = generar_word(datos)
    st.success("Checklist generado exitosamente.")
    st.download_button("📥 Descargar Word", archivo, file_name="Checklist_PROMTO_Receptor.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
