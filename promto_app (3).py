
import streamlit as st
from docx import Document
from datetime import datetime
import io

st.set_page_config(page_title="PROMTO App", layout="centered")

st.title("ğŸ“‹ PROMTO App â€“ Checklist Trasplante Renal")

# FunciÃ³n para generar el documento Word
def generar_checklist(receptor, donador, valoraciones, citas):
    doc = Document()
    doc.add_heading("Checklist PROMTO", level=1)

    # Receptor
    doc.add_paragraph(f"ğŸ§â€â™‚ï¸ Receptor: {receptor['nombre']}")
    for campo, valor in receptor['estudios'].items():
        doc.add_paragraph(f"â˜‘ {campo}  Fecha: {valor['fecha']}\n    Observaciones: {valor['obs']}")

    # Donador
    doc.add_paragraph(f"\nğŸ§â€â™€ï¸ Donador: {donador['nombre']}")
    doc.add_paragraph(f"â˜‘ IMC < 27  Fecha: {donador['imc_fecha']}\n    IMC: {donador['imc_valor']}")
    for campo, valor in donador['estudios'].items():
        doc.add_paragraph(f"â˜‘ {campo}  Fecha: {valor['fecha']}\n    Observaciones: {valor['obs']}")

    # Campos finales
    doc.add_paragraph("\nğŸ“ Valoraciones adicionales realizadas:")
    doc.add_paragraph(valoraciones)

    doc.add_paragraph("\nğŸ“… Citas pendientes por agendar:")
    doc.add_paragraph(citas)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# Captura de datos
with st.form("form_promto"):
    st.subheader("ğŸ§â€â™‚ï¸ Receptor")
    receptor_nombre = st.text_input("Nombre del receptor")
    receptor_estudios = {}
    for campo in ["Ecocardiograma", "TomografÃ­a de abdomen", "Exudado nasal"]:
        col1, col2 = st.columns(2)
        with col1:
            fecha = st.date_input(f"Fecha de {campo}", format="DD/MM/YYYY", key=f"{campo}_fecha")
        with col2:
            obs = st.text_input(f"Observaciones de {campo}", key=f"{campo}_obs")
        receptor_estudios[campo] = {"fecha": fecha.strftime("%d/%m/%Y"), "obs": obs}

    st.subheader("ğŸ§â€â™€ï¸ Donador")
    donador_nombre = st.text_input("Nombre del donador")
    imc_fecha = st.date_input("Fecha de cÃ¡lculo de IMC", format="DD/MM/YYYY")
    imc_valor = st.text_input("Valor de IMC")

    donador_estudios = {}
    for campo in ["Quantiferon", "SerologÃ­a VIH"]:
        col1, col2 = st.columns(2)
        with col1:
            fecha = st.date_input(f"Fecha de {campo}", format="DD/MM/YYYY", key=f"{campo}_fecha")
        with col2:
            obs = st.text_input(f"Observaciones de {campo}", key=f"{campo}_obs")
        donador_estudios[campo] = {"fecha": fecha.strftime("%d/%m/%Y"), "obs": obs}

    valoraciones_extra = st.text_area("Valoraciones adicionales realizadas")
    citas_pendientes = st.text_area("Citas pendientes por agendar")

    submitted = st.form_submit_button("Generar Checklist")

    if submitted:
        file_bytes = generar_checklist(
            {"nombre": receptor_nombre, "estudios": receptor_estudios},
            {"nombre": donador_nombre, "imc_fecha": imc_fecha.strftime("%d/%m/%Y"), "imc_valor": imc_valor, "estudios": donador_estudios},
            valoraciones_extra,
            citas_pendientes
        )
        st.success("âœ… Checklist generado. Descarga abajo:")
        st.download_button(
            "ğŸ“¥ Descargar Checklist Word",
            file_bytes,
            file_name=f"Checklist_PROMTO_{receptor_nombre.replace(' ', '_')}_{donador_nombre.replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
